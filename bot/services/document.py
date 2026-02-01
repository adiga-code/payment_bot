# bot/services/document.py

import logging
import os
from datetime import datetime, timedelta
from decimal import Decimal

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from database import DocumentRepository, ApplicationRepository
from utils.amount_words import amount_to_words

logger = logging.getLogger(__name__)

DOCUMENTS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), 'documents')

# Русские названия месяцев (родительный падеж)
_MONTHS_GEN = [
    '', 'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
    'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
]


def _format_amount(amount: Decimal) -> str:
    """100000.00 -> 100 000,00"""
    integer = int(amount)
    frac = int(round((amount - integer) * 100))
    int_str = f"{integer:,}".replace(",", " ")
    return f"{int_str},{frac:02d}"


def _format_date_ru(dt: datetime) -> str:
    """datetime -> '23 января 2026 г.'"""
    return f"{dt.day} {_MONTHS_GEN[dt.month]} {dt.year} г."


def _set_cell(cell, text: str, bold=False, size=8, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Установить текст ячейки с форматированием"""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Arial'


class DocumentService:
    """Сервис генерации документов (счетов на оплату)"""

    def __init__(self, document_repo: DocumentRepository, application_repo: ApplicationRepository):
        self.document_repo = document_repo
        self.application_repo = application_repo
        os.makedirs(DOCUMENTS_DIR, exist_ok=True)

    async def generate_invoice(self, app_id: int) -> str | None:
        """
        Генерирует счёт на оплату (DOCX) для одобренной заявки.

        Returns:
            file_path — путь к сгенерированному файлу, или None при ошибке
        """
        app = await self.application_repo.get_with_relations(app_id)
        if not app:
            logger.error(f"Application {app_id} not found for invoice generation")
            return None

        # Деактивируем старые счета этой заявки
        await self.document_repo.deactivate_old(app_id, 'invoice')

        # Создаём запись документа в БД (чтобы получить ID для номера)
        doc_record = await self.document_repo.create(
            application_id=app_id,
            type='invoice',
        )

        # Номер счёта
        invoice_number = str(doc_record.id)
        now = datetime.utcnow()

        # Данные компании-поставщика
        co = app.company
        company_name = co.name if co else "—"
        company_inn = co.inn if co else "—"
        company_kpp = co.kpp if co and co.kpp else ""
        company_address = co.legal_address if co and co.legal_address else ""
        company_bank_name = co.bank_name if co and co.bank_name else ""
        company_bik = co.bank_bik if co and co.bank_bik else ""
        company_account = co.bank_account if co and co.bank_account else ""
        company_corr = co.bank_corr_account if co and co.bank_corr_account else ""
        director_name = co.director_name if co and co.director_name else ""
        accountant_name_val = co.accountant_name if co and co.accountant_name else ""

        # Данные покупателя (плательщика)
        payer_name = app.payer_name or "—"
        payer_inn = app.payer_inn or "—"
        payer_kpp = app.payer_kpp or ""
        payer_address = app.payer_address or ""

        # Данные договора
        contract_number = app.contract_number or ""
        contract_date = app.contract_date or ""
        invoice_purpose = app.invoice_purpose or app.purpose or ""

        amount = app.amount
        amount_str = _format_amount(amount)

        # НДС 22%
        nds = amount * Decimal('22') / Decimal('122')
        nds_str = _format_amount(nds)

        # Назначение строки таблицы
        line_desc = f"Оплата по счету №{invoice_number} от {now.strftime('%d.%m.%Y')}"
        if contract_number:
            line_desc += f" к договору №{contract_number}"
            if contract_date:
                line_desc += f" от {contract_date}"
        if invoice_purpose:
            line_desc += f" за {invoice_purpose}"

        # Основание
        basis = ""
        if contract_number:
            basis = contract_number
            if contract_date:
                basis += f" от {contract_date}"

        # Дата оплаты — +3 дня от сегодня
        pay_by = now + timedelta(days=3)

        # Генерируем DOCX
        file_name = f"СЧ-{invoice_number}.docx"
        file_path = os.path.join(DOCUMENTS_DIR, file_name)

        try:
            self._build_invoice_docx(
                file_path=file_path,
                invoice_number=invoice_number,
                invoice_date=now,
                company_name=company_name,
                company_inn=company_inn,
                company_kpp=company_kpp,
                company_address=company_address,
                company_bank_name=company_bank_name,
                company_bik=company_bik,
                company_account=company_account,
                company_corr=company_corr,
                director_name=director_name,
                accountant_name_val=accountant_name_val,
                payer_name=payer_name,
                payer_inn=payer_inn,
                payer_kpp=payer_kpp,
                payer_address=payer_address,
                basis=basis,
                line_desc=line_desc,
                amount=amount,
                amount_str=amount_str,
                nds_str=nds_str,
                pay_by=pay_by,
            )
        except Exception as e:
            logger.error(f"Failed to generate invoice DOCX for app {app_id}: {e}")
            return None

        # Обновляем запись в БД
        display_number = f"СЧ-{invoice_number}"
        await self.document_repo.update_by_id(
            doc_record.id,
            number=display_number,
            file_path=file_path,
        )

        # Обновляем статус заявки
        await self.application_repo.mark_invoice_generated(app_id)

        logger.info(f"Invoice {display_number} generated for application {app.external_id}")
        return file_path

    def _build_invoice_docx(
        self,
        file_path: str,
        invoice_number: str,
        invoice_date: datetime,
        company_name: str,
        company_inn: str,
        company_kpp: str,
        company_address: str,
        company_bank_name: str,
        company_bik: str,
        company_account: str,
        company_corr: str,
        director_name: str,
        accountant_name_val: str,
        payer_name: str,
        payer_inn: str,
        payer_kpp: str,
        payer_address: str,
        basis: str,
        line_desc: str,
        amount: Decimal,
        amount_str: str,
        nds_str: str,
        pay_by: datetime,
    ):
        """Создаёт DOCX-файл счёта на оплату (полный формат)"""
        doc = DocxDocument()

        # Стиль
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(9)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)

        # Узкие поля
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(1.5)

        # ==================== ШАПКА БАНКА ====================
        bank_table = doc.add_table(rows=5, cols=4)
        bank_table.style = 'Table Grid'

        # Строка 1: Банк получателя | БИК
        _set_cell(bank_table.rows[0].cells[0], company_bank_name, size=8)
        _set_cell(bank_table.rows[0].cells[2], 'БИК', bold=True, size=8)
        _set_cell(bank_table.rows[0].cells[3], company_bik, size=8)

        # Строка 2: (продолжение банка) | Сч.№ (корр.)
        _set_cell(bank_table.rows[1].cells[0], 'Банк получателя', size=7)
        _set_cell(bank_table.rows[1].cells[2], 'Сч. №', bold=True, size=8)
        _set_cell(bank_table.rows[1].cells[3], company_corr, size=8)

        # Строка 3: ИНН | КПП | Сч.№ (расчётный)
        inn_kpp = f'ИНН {company_inn}'
        if company_kpp:
            inn_kpp += f'     КПП {company_kpp}'
        _set_cell(bank_table.rows[2].cells[0], inn_kpp, size=8)
        _set_cell(bank_table.rows[2].cells[2], 'Сч. №', bold=True, size=8)
        _set_cell(bank_table.rows[2].cells[3], company_account, size=8)

        # Строка 4: Название компании
        _set_cell(bank_table.rows[3].cells[0], company_name, bold=True, size=9)

        # Строка 5: Получатель
        _set_cell(bank_table.rows[4].cells[0], 'Получатель', size=7)

        # Объединяем ячейки где нужно
        bank_table.rows[0].cells[0].merge(bank_table.rows[0].cells[1])
        bank_table.rows[1].cells[0].merge(bank_table.rows[1].cells[1])
        bank_table.rows[2].cells[0].merge(bank_table.rows[2].cells[1])
        bank_table.rows[3].cells[0].merge(bank_table.rows[3].cells[1])
        bank_table.rows[3].cells[2].merge(bank_table.rows[3].cells[3])
        bank_table.rows[4].cells[0].merge(bank_table.rows[4].cells[1])
        bank_table.rows[4].cells[2].merge(bank_table.rows[4].cells[3])

        doc.add_paragraph()

        # ==================== ЗАГОЛОВОК СЧЁТА ====================
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = title.add_run(
            f'Счёт на оплату № {invoice_number} от {_format_date_ru(invoice_date)}'
        )
        run.bold = True
        run.font.size = Pt(13)

        # Горизонтальная линия (через нижнюю границу параграфа)
        doc.add_paragraph('─' * 80)

        # ==================== ПОСТАВЩИК / ПОКУПАТЕЛЬ ====================
        supplier_line = f'{company_name}, ИНН {company_inn}'
        if company_kpp:
            supplier_line += f', КПП {company_kpp}'
        if company_address:
            supplier_line += f', {company_address}'

        p = doc.add_paragraph()
        run = p.add_run('Поставщик\n(Исполнитель):  ')
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(supplier_line)
        run.font.size = Pt(9)

        buyer_line = f'{payer_name}, ИНН {payer_inn}'
        if payer_kpp:
            buyer_line += f', КПП {payer_kpp}'
        if payer_address:
            buyer_line += f', {payer_address}'

        p = doc.add_paragraph()
        run = p.add_run('Покупатель\n(Заказчик):  ')
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(buyer_line)
        run.font.size = Pt(9)

        if basis:
            p = doc.add_paragraph()
            run = p.add_run('Основание:  ')
            run.bold = True
            run.font.size = Pt(9)
            run = p.add_run(basis)
            run.font.size = Pt(9)

        doc.add_paragraph()

        # ==================== ТАБЛИЦА ТОВАРОВ ====================
        tbl = doc.add_table(rows=2, cols=6)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = ['№', 'Товары (работы, услуги)', 'Кол-во', 'Ед.', 'Цена', 'Сумма']
        for i, h in enumerate(headers):
            _set_cell(tbl.rows[0].cells[i], h, bold=True, size=8,
                       align=WD_ALIGN_PARAGRAPH.CENTER)

        # Данные строки
        _set_cell(tbl.rows[1].cells[0], '1', size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(tbl.rows[1].cells[1], line_desc, size=8)
        _set_cell(tbl.rows[1].cells[2], '1', size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(tbl.rows[1].cells[3], 'шт', size=8, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell(tbl.rows[1].cells[4], amount_str, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell(tbl.rows[1].cells[5], amount_str, size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)

        doc.add_paragraph()

        # ==================== ИТОГО ====================
        for label, value in [
            ('Итого:', amount_str),
            ('В том числе НДС 22%:', nds_str),
            ('Всего к оплате:', amount_str),
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f'{label}  ')
            run.font.size = Pt(9)
            run = p.add_run(value)
            run.bold = True
            run.font.size = Pt(9)

        doc.add_paragraph()

        # ==================== СУММА ПРОПИСЬЮ ====================
        words = amount_to_words(amount)
        p = doc.add_paragraph()
        run = p.add_run(f'Всего наименований 1, на сумму {amount_str} руб.')
        run.font.size = Pt(9)
        p = doc.add_paragraph()
        run = p.add_run(words)
        run.bold = True
        run.font.size = Pt(9)

        doc.add_paragraph()

        # ==================== УСЛОВИЯ ОПЛАТЫ ====================
        pay_by_str = pay_by.strftime('%d.%m.%Y')
        terms = (
            f'Оплатить не позднее {pay_by_str}\n'
            f'Оплата данного счёта означает согласие с условиями поставки товара.\n'
            f'Уведомление об оплате обязательно, в противном случае не гарантируется наличие товара на складе.\n'
            f'Товар отпускается по факту прихода денег на р/с Поставщика, самовывозом, при наличии доверенности и паспорта.'
        )
        p = doc.add_paragraph()
        run = p.add_run(terms)
        run.font.size = Pt(8)

        doc.add_paragraph()
        doc.add_paragraph('─' * 80)

        # ==================== ПОДПИСИ ====================
        p = doc.add_paragraph()
        run = p.add_run('Руководитель  ____________  ')
        run.font.size = Pt(9)
        run = p.add_run(director_name)
        run.bold = True
        run.font.size = Pt(9)

        run = p.add_run('          Бухгалтер  ____________  ')
        run.font.size = Pt(9)
        run = p.add_run(accountant_name_val)
        run.bold = True
        run.font.size = Pt(9)

        doc.save(file_path)
